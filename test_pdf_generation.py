# test_pdf_generation.py
import os
import django
from docx import Document

# --- 1. Configurar o Ambiente Django ---
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'loja.settings')
django.setup()

# --- 2. Importar os Módulos Necessários ---
from quoteflow.models import Cotacao
from quoteflow.utils import criar_replacements, converter_docx_para_pdf
from django.contrib.auth import get_user_model

# ==============================================================================
# --- FUNÇÃO DE TESTE ROBUSTA (lida com tags divididas em vários 'runs') ---
# ==============================================================================

def substituir_tags_documento_robusto(doc, replacements):
    """
    Função de substituição robusta que encontra e substitui texto que pode
    estar dividido em múltiplos 'runs' dentro de um parágrafo.
    """
    for key, value in replacements.items():
        key = str(key)
        value = str(value)

        # Procura nos parágrafos do corpo principal
        for p in doc.paragraphs:
            replace_text_in_paragraph(p, key, value)
        
        # Procura nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_text_in_paragraph(p, key, value)
        
        # Procura nos cabeçalhos e rodapés
        for section in doc.sections:
            for p in section.header.paragraphs:
                replace_text_in_paragraph(p, key, value)
            for p in section.footer.paragraphs:
                replace_text_in_paragraph(p, key, value)

def replace_text_in_paragraph(paragraph, key, value):
    """
    Função auxiliar que efetivamente substitui o texto, lidando com 'runs'.
    Esta função é chamada repetidamente para cada par de chave/valor.
    """
    # Concatena o texto de todos os runs para encontrar a chave
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return # Se a chave não está no parágrafo, não faz nada

    # Encontrou a chave, agora vamos fazer a substituição
    # Primeiro, substituímos a primeira ocorrência da chave no primeiro run que a contém
    # e limpamos o resto da chave nos runs seguintes.
    
    run_texts = [run.text for run in paragraph.runs]
    # Limpa todos os runs para reescrever
    for run in paragraph.runs:
        run.text = ""
    
    # Reescreve o texto com a substituição, no primeiro run
    paragraph.runs[0].text = full_text.replace(key, value, 1)

# ==============================================================================
# --- LÓGICA PRINCIPAL DO SCRIPT DE TESTE ---
# ==============================================================================

print("Ambiente Django configurado. Iniciando teste com SCRIPT DE SUBSTITUIÇÃO ROBUSTO...")

# --- Definir os Parâmetros do Teste ---
COTACAO_ID_PARA_TESTE = 10571
USER_ID_PARA_TESTE = 16

OUTPUT_TEST_DIR = os.path.join(os.path.dirname(__file__), 'test_outputs')
os.makedirs(OUTPUT_TEST_DIR, exist_ok=True)
print(f"Os arquivos de resultado serão salvos em: {OUTPUT_TEST_DIR}")

try:
    # --- Preparar os Dados ---
    print(f"Buscando Cotação ID: {COTACAO_ID_PARA_TESTE} e Usuário ID: {USER_ID_PARA_TESTE}")
    cotacao = Cotacao.objects.get(pk=COTACAO_ID_PARA_TESTE)
    User = get_user_model()
    user = User.objects.get(pk=USER_ID_PARA_TESTE)
    
    empresa_usuario = user.perfil.empresa
    template_path = empresa_usuario.template_proposta.path
    print(f"Usando template: {template_path}")

    # Atenção: o número da proposta aqui usa o gerado pela função original
    # para que a tag #PROPR#REV seja substituída corretamente.
    from quoteflow.utils import gerar_numero_proposta
    numero_proposta = gerar_numero_proposta(cotacao)
    replacements = criar_replacements(cotacao, numero_proposta, user)

    # --- Etapa 1: Processar o Documento Word com a LÓGICA ROBUSTA ---
    print("\n--- Etapa 1: Processando o template .docx com a FUNÇÃO ROBUSTA ---")
    doc = Document(template_path)
    
    substituir_tags_documento_robusto(doc, replacements)

    output_docx_filename = f"Proposta_{numero_proposta}_INTERMEDIARIO_FINAL.docx"
    output_docx_path = os.path.join(OUTPUT_TEST_DIR, output_docx_filename)
    
    doc.save(output_docx_path)
    print(f"✅ Arquivo .docx intermediário salvo em: {output_docx_path}")
    print("   -> Abra este arquivo para verificar se TODAS as tags foram substituídas.")

    # --- Etapa 2: Converter para PDF ---
    print("\n--- Etapa 2: Convertendo o .docx para .pdf ---")
    final_pdf_path = converter_docx_para_pdf(output_docx_path, output_dir=OUTPUT_TEST_DIR)
    
    print("-" * 50)
    print("✅ SUCESSO FINAL!")
    print(f"Proposta em PDF gerada com sucesso em: {final_pdf_path}")
    print("Verifique se o PDF final está 100% correto.")
    print("-" * 50)

except Exception as e:
    print(f"Ocorreu um erro inesperado: {e}")