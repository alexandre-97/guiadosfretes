from .models import Cotacao
from urllib.parse import quote
import requests
from datetime import datetime
import os
import re
import unicodedata
from docx import Document
from django.template.loader import render_to_string
from django.core import mail
import smtplib
import ssl
from django.conf import settings
from django.core.mail import EmailMultiAlternatives
from .models import EmailEnviado
import subprocess
import base64
import logging
from django.contrib.auth import get_user_model
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from perfil.models import ConfiguracaoEmail
from django.core.files.base import ContentFile
import random
import time
from requests.exceptions import RequestException, SSLError, ProxyError 
from requests.adapters import HTTPAdapter
from urllib3.util.ssl_ import create_urllib3_context
import smtplib
import ssl
from django.core.mail import EmailMultiAlternatives
logger = logging.getLogger(__name__)
from .models import EmailEnviado

SPINTAX_VARIATIONS = {
    "saudacao_hora": {
        "manha": ["Bom dia", "Tenha um ótimo dia"],
        "tarde": ["Boa tarde", "Tenha uma ótima tarde"],
        "noite": ["Boa noite", "Tenha uma boa noite"]
    },
    "proposta": {
        "saudacao_contato": ["👋 Olá, {contato}!", "Olá, {contato}, tudo bem?", "Prezado(a) {contato},"],
        "introducao": [
            "Segue a proposta da *{empresa}* que você solicitou.",
            "Conforme solicitado, aqui está sua cotação da *{empresa}*.",
            "Recebemos seu pedido de cotação. Confira abaixo a proposta da *{empresa}*:"
        ],
        "detalhes_proposta": ["📄 Proposta:", "📄 Cotação Nº:", "📄 Cotação:"],
        "detalhes_origem": ["🚚 Origem:", "📍 Coleta em:", "🏳️ De:"],
        "detalhes_destino": ["🚛 Destino:", "🏁 Entrega em:", "➡️ Para:"],
        "detalhes_valor": ["💰 Valor:", "💲 Total:", "💵 Frete:"],
        "detalhes_coleta": ["⏱️ Coleta:", "🗓️ Prazo Coleta:"],
        "detalhes_entrega": ["🏁 Entrega:", "🗓️ Prazo Entrega:"],
        "detalhes_pagamento": ["💵 Pagamento:", "💵 Forma de Pgto:"],
        "detalhes_seguro": ["🔒 Seguro:", "🛡️ Cobertura:"],
        "negociacao": [
            "Estamos abertos a negociações. Se precisar ajustar preço, prazo ou forma de pagamento, será um prazer lhe atender.",
            "Caso precise de alguma alteração no valor ou prazo, por favor, me avise. Estamos à disposição para negociar.",
            "Se esta proposta não atender completamente sua necessidade, me diga o que podemos fazer para ajustá-la. 😉"
        ]
    },
    "proposta_v2": {
        "apresentacao": [
            "Aqui está a sua proposta de frete da *{empresa}*!",
            "Obrigado por cotar conosco! Confira os detalhes da sua proposta:",
            "Sua cotação com a *{empresa}* está pronta! Veja abaixo:"
        ],
        "chamada_negociacao": [
            "Algum detalhe precisa de ajuste? Me avise!",
            "Se os valores ou prazos não baterem, me chame para negociarmos.",
            "Podemos conversar sobre os detalhes. O que acha?"
        ]
    },
    "solicitacao_medidas": {
        "saudacao": ["{contato}, tudo bem?", "Olá, {contato}!", "Oi, {contato}, como vai?"],
        "pedido": [
            "Para calcularmos o frete corretamente para a proposta *{numero_proposta}*, precisamos das *medidas* da sua carga. Poderia nos informar abaixo?",
            "Referente à sua cotação *{numero_proposta}*, por favor, nos envie as *dimensões (Comprimento x Largura x Altura)* de cada volume.",
            "Para finalizarmos sua cotação *{numero_proposta}*, só falta nos informar as *medidas*. Poderia preencher abaixo, por favor?"
        ],
        "instrucao_geral": [
            "Para cargas gerais, por favor, informe as *medidas de cada volume* (ex: 1,20m x 0,80m x 1,00m).",
            "Precisamos das dimensões para calcular a cubagem. Por favor, envie o *comprimento, a largura e a altura* de cada item.",
            "Por favor, detalhe as *medidas individuais* de cada volume para que possamos dar continuidade."
        ],
        "despedida": ["Aguardamos seu retorno para finalizar a proposta.", "Fico no aguardo das informações.", "Obrigado(a) pela ajuda!"]
    },
    "solicitacao_coleta": {
        "introducao": [
            "Olá! Para agendar a coleta da proposta *{numero_proposta}*, precisamos que preencha os dados abaixo.",
            "Tudo certo para a coleta da sua carga! Referente à proposta *{numero_proposta}*, por favor, confirme as informações a seguir.",
            "Segue o formulário para a coleta da proposta *{numero_proposta}*."
        ],
        "pedido": ["Por favor, preencha os campos em branco.", "Complete as informações abaixo para agendarmos.", "Precisamos desses dados para a programação da coleta."]
    },
    "solicitacao_pagamento": {
        "saudacao": ["😃 Olá {cliente},", "Oi {cliente}, temos boas notícias!", "Prezado(a) {cliente}, tudo bem?"],
        "noticia_boa": ["boas notícias!", "uma ótima notícia sobre sua entrega!", "uma atualização sobre sua carga."],
        "status_mercadoria": [
            "Sua mercadoria foi coletada com sucesso e já está a caminho de nossa base.",
            "Informamos que sua carga foi retirada pelo nosso motorista e segue para nossa unidade.",
            "Coleta realizada! Sua mercadoria está em trânsito para nosso centro de distribuição."
        ],
        "pedido_comprovante": [
            "Se possível, nos envie o comprovante de pagamento para agilizarmos a liberação da entrega.",
            "Assim que realizar a transferência, por gentileza, nos encaminhe o comprovante.",
            "O envio do comprovante nos ajuda a identificar o pagamento e acelerar o processo. Agradecemos!"
        ],
        "despedida": ["Desde já agradecemos a parceria.", "Qualquer dúvida, estamos à disposição.", "Agradecemos pela confiança!"]
    }
}
class TlsAdapter(HTTPAdapter):
    def init_poolmanager(self, *args, **kwargs):
        context = create_urllib3_context(
            ciphers='DEFAULT:@SECLEVEL=2',
            cert_reqs=ssl.CERT_REQUIRED,
            options=ssl.OP_NO_SSLv2 | ssl.OP_NO_SSLv3 | ssl.OP_NO_TLSv1 | ssl.OP_NO_TLSv1_1
        )
        kwargs['ssl_context'] = context
        return super(TlsAdapter, self).init_poolmanager(*args, **kwargs)
def _get_secure_requests_session():
    """Cria e retorna uma sessão de 'requests' com o adaptador de segurança customizado."""
    session = requests.Session()
    session.mount('https://', TlsAdapter())
    return session

def log_action(logger_instance, level, user, message, cotacao=None):
    """
    Formata e registra uma ação do usuário de forma padronizada.
    
    Args:
        logger_instance: A instância do logger (ex: logger).
        level (str): O nível do log ('info', 'warning', 'error', 'debug').
        user: O objeto User do Django.
        message (str): A mensagem principal do log.
        cotacao (Cotacao, optional): O objeto de cotação relacionado.
    """
    try:
        empresa_nome = "N/A"
        if hasattr(user, 'perfil') and user.perfil and user.perfil.empresa:
            empresa_nome = user.perfil.empresa.nome

        proposta_id = f"ID:{cotacao.id}" if cotacao else "N/A"
        if cotacao and hasattr(cotacao, 'proposta_id_url'):
            proposta_id = cotacao.proposta_id_url

        log_prefix = f"[USER: {user.username} | EMPRESA: {empresa_nome} | COTAÇÃO: {proposta_id}]"
        
        log_func = getattr(logger_instance, level)
        log_func(f"{log_prefix} {message}")

    except Exception:
        # Fallback para o caso de algum objeto ser inválido
        logger_instance.error(f"Falha ao formatar log. Mensagem original: {message}")


class WhatsAppNumberInvalidException(Exception):
    """Exceção customizada para erros de número de WhatsApp inválido."""
    pass
class WhatsAppInstanceDisconnectedException(Exception):
    """Exceção para quando a instância do WhatsApp não está logada."""
    pass

def valor_por_extenso(valor):
    unidades = ['zero', 'um', 'dois', 'três', 'quatro', 'cinco', 'seis', 'sete', 'oito', 'nove']
    dez_a_dezenove = ['dez', 'onze', 'doze', 'treze', 'quatorze', 'quinze', 'dezesseis', 'dezessete', 'dezoito', 'dezenove']
    dezenas = ['', '', 'vinte', 'trinta', 'quarenta', 'cinquenta', 'sessenta', 'setenta', 'oitenta', 'noventa']
    centenas = ['', 'cento', 'duzentos', 'trezentos', 'quatrocentos', 'quinhentos', 'seiscentos', 'setecentos', 'oitocentos', 'novecentos']

    def extenso_numerico(n):
        if n == 0:
            return ''
        elif n < 10:
            return unidades[n]
        elif n < 20:
            return dez_a_dezenove[n - 10]
        elif n < 100:
            dezena = n // 10
            unidade = n % 10
            return dezenas[dezena] + (f' e {unidades[unidade]}' if unidade else '')
        elif n < 1000:
            if n == 100:
                return 'cem'
            centena = n // 100
            resto = n % 100
            return centenas[centena] + (f' e {extenso_numerico(resto)}' if resto else '')
        elif n < 1000000:
            milhar = n // 1000
            resto = n % 1000
            if milhar == 1:
                prefixo = 'mil'
            else:
                prefixo = f'{extenso_numerico(milhar)} mil'
            return prefixo + (f' e {extenso_numerico(resto)}' if resto else '')
        else:
            return 'valor muito alto'

    try:
        valor = round(float(valor), 2)
        reais = int(valor)
        centavos = int(round((valor - reais) * 100))

        partes = []

        if reais > 0:
            partes.append(extenso_numerico(reais))
            partes.append('real' if reais == 1 else 'reais')

        if centavos > 0:
            if reais > 0:
                partes.append('e')
            partes.append(extenso_numerico(centavos))
            partes.append('centavo' if centavos == 1 else 'centavos')

        return ' '.join(partes) if partes else 'zero real'

    except (ValueError, TypeError):
        return "valor inválido"

def sanitize_text(text):
    """Sanitiza texto para evitar problemas de encoding."""
    if text is None:
        return ''
    
    try:
        if isinstance(text, bytes):
            text = text.decode('utf-8', errors='replace')
        
        text = text.encode('utf-8', errors='replace').decode('utf-8')
        text = ''.join(char for char in text if ord(char) >= 32 or ord(char) == 10 or ord(char) == 13)
        
        return text
    except Exception:
        return str(text)[:500]

def gerar_numero_proposta(cotacao):
    """Gera o número da proposta usando o número sequencial da empresa."""
    try:
        # Garante que a cotação e a empresa existem
        if not cotacao or not cotacao.empresa:
            raise ValueError("Cotação ou empresa associada não encontrada.")
            
        empresa = cotacao.empresa
        nome_empresa = sanitize_text(empresa.nome)[:3].upper()
        
        # Usa o novo campo sequencial. Se ele não existir, usa o ID como fallback.
        sequencial = cotacao.numero_sequencial_empresa or cotacao.id
        
        return f"{nome_empresa}{str(sequencial).zfill(3)}QF"
    except Exception as e:
        # Se algo der errado, usa o ID como fallback e loga um erro
        logger.error(f"Erro ao gerar número de proposta para cotação {getattr(cotacao, 'id', 'N/A')}: {e}. Usando ID como fallback.")
        nome_empresa = "EMP"
        if hasattr(cotacao, 'empresa') and cotacao.empresa:
            nome_empresa = sanitize_text(cotacao.empresa.nome)[:3].upper()
        return f"{nome_empresa}{str(getattr(cotacao, 'id', '0')).zfill(3)}QF"
    

def criar_especificacao_carga(cotacao):
    """Cria a string de especificação da carga formatada."""
    valor_mercadoria = f"R$ {float(cotacao.valor_mercadoria):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if cotacao.valor_mercadoria else 'R$ 0,00'
    
    return "\n".join([
        f"Nº VOLUMES: {sanitize_text(cotacao.volumes) or 'NÃO INFORMADO'}",
        f"CUBAGEM: {sanitize_text(cotacao.cubagem) or '0,00'} M³",
        f"PESO: {sanitize_text(cotacao.peso) or 'NÃO INFORMADO'} KG",
        f"VALOR NF: {valor_mercadoria}",
        "",
        f"OBSERVAÇÃO:\n{sanitize_text(cotacao.observacao) or 'NENHUMA OBSERVAÇÃO'}"
    ])

def criar_replacements(cotacao, numero_proposta, user):
    """Cria o dicionário de substituições para o documento."""
    log_action(logger, 'debug', user, "Iniciando criação de replacements para documento.", cotacao)

    valor_mercadoria = f"R$ {float(cotacao.valor_mercadoria):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if cotacao.valor_mercadoria else 'R$ 0,00'
    valor_proposta = f"R$ {float(cotacao.valor_proposta):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if cotacao.valor_proposta else 'R$ 0,00'
    
    # Busca o número de WhatsApp do perfil do usuário logado de forma segura
    whatsapp_usuario_bruto = None
    print(f"User object recebido: {user}") # DEBUG
    if hasattr(user, 'perfil') and user.perfil:
        print(f"Perfil do usuário encontrado: {user.perfil}") # DEBUG
        whatsapp_usuario_bruto = user.perfil.telefone_whatsapp
    else:
        print("AVISO: Usuário não tem perfil (user.perfil) associado.") # DEBUG

    # Formata o número usando a nova função
    whatsapp_usuario_formatado = formatar_whatsapp_para_exibicao(whatsapp_usuario_bruto)
    
    replacements = {
        '#CONTATO': sanitize_text(cotacao.contato or 'Cliente'),
        '#PROPR#REV': numero_proposta,
        '#EMAIL': sanitize_text(cotacao.email),
        '#TELEFONE': sanitize_text(cotacao.telefone or 'NÃO INFORMADO'), # Telefone da Cotação
        '#DATA1': datetime.now().strftime('%d/%m/%Y'),
        '#ESPECIFICACAO': criar_especificacao_carga(cotacao),
        '#COLETA': sanitize_text(cotacao.origem or 'NÃO INFORMADO'),
        '#DESTINO': sanitize_text(cotacao.destino or 'NÃO INFORMADO'),
        '#PRAZOCOLETA': sanitize_text(cotacao.prazo_coleta or 'A COMBINAR'),
        '#PRAZOENTREGA': sanitize_text(cotacao.prazo_entrega or 'A COMBINAR'),
        '#VALOR': valor_proposta,
        '#EXTENSO': valor_por_extenso(cotacao.valor_proposta),
        '#TXT1': '',
        '#TXT2': '',
        '#TXT3': '',
        '#TXT4': '',
        '#ATEN': sanitize_text(cotacao.contato or 'Cliente'),
        '#ESPECIFICACAONº VOLUMES': f"Nº VOLUMES: {sanitize_text(cotacao.volumes) or 'NÃO INFORMADO'}",
        '#VALOR_NF': valor_mercadoria,
        
        # --- TAG DO USUÁRIO ---
        '#WHATSAPP': whatsapp_usuario_formatado,
    }
    return replacements

def _replace_text_in_paragraph(paragraph, key, value):
    """
    Função auxiliar que efetivamente substitui o texto, lidando com tags
    divididas em múltiplos 'runs' (pedaços de texto formatado).
    """
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return

    new_text = full_text.replace(key, value)
    
    if paragraph.runs:
        # Apaga o conteúdo de todos os runs para evitar duplicar texto
        for run in paragraph.runs:
            run.text = ""
        
        # Escreve o novo texto completo no primeiro run, que manterá a formatação
        paragraph.runs[0].text = new_text

def substituir_tags_documento(doc, replacements):
    """
    Substitui as tags de um documento DOCX de forma robusta e não-destrutiva,
    preservando a formatação de fontes, ícones e estilos.
    """
    for key, value in replacements.items():
        key_str = str(key)
        value_str = str(value)

        # 1. Procura e substitui nos parágrafos do corpo principal
        for p in doc.paragraphs:
            _replace_text_in_paragraph(p, key_str, value_str)
        
        # 2. Procura e substitui nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_text_in_paragraph(p, key_str, value_str)
        
        # 3. Procura e substitui nos cabeçalhos e rodapés
        # O código original tinha uma recursão aqui que causava o erro
        # '_Header' object has no attribute 'sections'. Esta é a forma correta.
        for section in doc.sections:
            for p in section.header.paragraphs:
                _replace_text_in_paragraph(p, key_str, value_str)
            for p in section.footer.paragraphs:
                _replace_text_in_paragraph(p, key_str, value_str)


def gerar_proposta_word(cotacao, replacements, empresa_usuario, template_path, user):
    """Gera o documento Word da proposta, usando o template especificado."""
    log_action(logger, 'info', user, "Iniciando geração de proposta em Word/PDF...", cotacao)
    
    if not os.path.exists(template_path):
        log_action(logger, 'error', user, f"Template de proposta não encontrado: {template_path}", cotacao)
        raise Exception("Template para proposta inexistente.")

    try:
        doc = Document(template_path)
        
        numero_proposta = gerar_numero_proposta(cotacao)
        output_filename = f"Proposta_{numero_proposta}_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(settings.MEDIA_ROOT, 'propostas', output_filename)
        
        if os.path.exists(output_path):
            os.remove(output_path)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        substituir_tags_documento(doc, replacements)
        
        doc.save(output_path)

        output_pdf_path = converter_docx_para_pdf(output_path)
        output_pdf_filename = os.path.basename(output_pdf_path)

        log_action(logger, 'info', user, f"Proposta gerada com sucesso: {output_pdf_filename}", cotacao)
        return output_pdf_path, output_pdf_filename, numero_proposta
        
    except Exception as e:
        log_action(logger, 'error', user, f"Erro ao gerar proposta Word: {str(e)}", cotacao)
        raise Exception(f'Erro ao gerar proposta Word: {str(e)}')


def obter_assinatura_digital(user, empresa):
    """
    Obtém a assinatura digital. Prioriza a assinatura personalizada do perfil do usuário
    e, como fallback, monta uma assinatura com os dados da empresa correta.
    """
    # 1. Tenta obter a assinatura personalizada do perfil do usuário
    try:
        # Garante que 'user' e 'user.perfil' existem antes de tentar usar
        if user and hasattr(user, 'perfil') and user.perfil and hasattr(user.perfil, 'assinatura_digital'):
            assinatura_pessoal = user.perfil.assinatura_digital()
            # Se o método retornar uma assinatura válida, a usamos
            if assinatura_pessoal:
                return assinatura_pessoal
    except Exception as e:
        logger.warning(f"Não foi possível obter assinatura personalizada para o usuário {user}: {e}")

    # 2. FALLBACK CORRIGIDO: Monta a assinatura com os dados da EMPRESA.
    #    Executado se o usuário não tiver assinatura pessoal ou se ocorrer um erro.
    if empresa:
        # ATENÇÃO: Verifique se os nomes dos campos (ex: 'email', 'telefone', 'site')
        # correspondem exatamente aos campos do seu modelo 'perfil.Empresa'.
        nome_empresa = empresa.nome
        email_empresa = getattr(empresa, 'email', '')      # Busca o email da empresa
        tel_empresa = getattr(empresa, 'telefone', '')    # Busca o telefone da empresa
        site_empresa = getattr(empresa, 'site', '')        # Busca o site da empresa

        # Monta a assinatura dinamicamente com os dados encontrados
        assinatura_fallback = f"Equipe Comercial - {nome_empresa}"
        
        contatos = []
        if email_empresa:
            contatos.append(f"Email: {email_empresa}")
        if tel_empresa:
            contatos.append(f"Tel: {tel_empresa}")
        
        # Adiciona a linha de contatos apenas se houver algum
        if contatos:
             assinatura_fallback += "\n" + " | ".join(contatos)
        
        # Adiciona o site se existir
        if site_empresa:
            assinatura_fallback += f"\n{site_empresa}"
            
        return assinatura_fallback

    # 3. Fallback final: Caso não haja nem usuário com perfil nem empresa.
    return "Equipe Comercial"

import smtplib
import ssl
from django.core.mail import EmailMultiAlternatives

# ... (outras funções) ...

def enviar_email_proposta(config_email, cotacao, output_path, output_filename, numero_proposta, empresa_usuario, request_user=None):
    try:
        assinatura_digital = obter_assinatura_digital(request_user, empresa_usuario)
        data_envio = datetime.now().strftime('%d/%m/%Y %H:%M')

        email_subject = sanitize_text(f"Proposta {numero_proposta} - {empresa_usuario.nome_full}")
        text_content = sanitize_text(f"Segue em anexo a proposta {numero_proposta}")
        
        html_content = render_to_string('cotacoes/email_proposta.html', {
            'cotacao': cotacao,
            'empresa': empresa_usuario,
            'config_email': config_email,
            'numero_proposta': numero_proposta,
            'especificacao_carga': criar_especificacao_carga(cotacao),
            'assinatura_digital': assinatura_digital,
            'data_envio': data_envio
        })
        html_content = sanitize_text(html_content)

        email_registro = EmailEnviado.objects.create(
            cotacao=cotacao,
            enviado_por=request_user,
            destinatario=cotacao.email,
            assunto=email_subject,
            corpo=html_content,
            enviado_com_sucesso=False
        )

        msg = EmailMultiAlternatives(
            subject=email_subject,
            body=text_content,
            from_email=config_email.email,
            to=[cotacao.email],
            reply_to=[config_email.email_resposta or config_email.email]
        )
        msg.attach_alternative(html_content, "text/html")
        
        with open(output_path, 'rb') as f:
            file_content = f.read()
            msg.attach(
                filename=output_filename,
                content=file_content,
                mimetype='application/pdf'
            )
            email_registro.anexo.save(output_filename, ContentFile(file_content), save=True)

        try:
            server = None
            if config_email.porta_smtp == 465:
                logger.info(f"Conectando via SMTP_SSL na porta 465 para {config_email.email}")
                context = ssl.create_default_context()
                server = smtplib.SMTP_SSL(config_email.servidor_smtp, config_email.porta_smtp, context=context, timeout=20)
                server.login(config_email.email, config_email.senha)
                # CORREÇÃO: Usar sendmail em vez de send_message
                server.sendmail(msg.from_email, msg.recipients(), msg.message().as_bytes())

            elif config_email.usar_tls_smtp:
                logger.info(f"Conectando via SMTP com STARTTLS na porta {config_email.porta_smtp} para {config_email.email}")
                server = smtplib.SMTP(config_email.servidor_smtp, config_email.porta_smtp, timeout=20)
                server.ehlo()
                server.starttls()
                server.ehlo()
                server.login(config_email.email, config_email.senha)
                # CORREÇÃO: Usar sendmail em vez de send_message
                server.sendmail(msg.from_email, msg.recipients(), msg.message().as_bytes())

            else: # Fallback para conexões não seguras
                logger.info(f"Conectando via SMTP padrão na porta {config_email.porta_smtp} para {config_email.email}")
                server = smtplib.SMTP(config_email.servidor_smtp, config_email.porta_smtp, timeout=20)
                server.login(config_email.email, config_email.senha)
                # CORREÇÃO: Usar sendmail em vez de send_message
                server.sendmail(msg.from_email, msg.recipients(), msg.message().as_bytes())

            # Finaliza a conexão se ela foi estabelecida
            if server:
                server.quit()

            salvar_email_enviado(config_email, msg)
            email_registro.enviado_com_sucesso = True
            email_registro.save()
            logger.info(f"Email para {cotacao.email} enviado com sucesso.")
            return msg

        except Exception as e:
            email_registro.erro = str(e)
            email_registro.save()
            logger.error(f"Falha ao enviar e-mail para {cotacao.email} via SMTP: {e}", exc_info=True)
            raise

    except Exception as e:
        logger.error(f"Erro geral ao preparar e-mail para cotação {cotacao.id}: {e}", exc_info=True)
        if 'email_registro' in locals() and not email_registro.erro:
            email_registro.erro = str(e)
            email_registro.save()
        raise
def enviar_email_simples(config_email, cotacao, assunto, mensagem, request_user=None, is_html=False):
    try:
        empresa = getattr(request_user, 'perfil', None) and request_user.perfil.empresa or cotacao.empresa
        data_envio = datetime.now().strftime('%d/%m/%Y %H:%M')
        rodape = f"\n\nEnviado em: {data_envio}"
        
        assunto = sanitize_text(assunto)

        # --- INÍCIO DA CORREÇÃO ---
        # Agora, a função verifica se a mensagem já é HTML
        if is_html:
            # Se for HTML, apenas adicionamos o rodapé e definimos um texto simples como fallback
            corpo_html = mensagem + rodape.replace('\n', '<br>')
            corpo_texto = "Para visualizar este e-mail, por favor, utilize um leitor de e-mail compatível com HTML."
        else:
            # Se for texto simples, fazemos como antes
            assinatura_digital = obter_assinatura_digital(request_user, empresa)
            rodape_completo = f"\n\n{assinatura_digital}{rodape}"
            corpo_html = sanitize_text(mensagem).replace('\n', '<br>') + rodape_completo.replace('\n', '<br>')
            corpo_texto = sanitize_text(mensagem) + rodape_completo
        # --- FIM DA CORREÇÃO ---

        email_registro = EmailEnviado.objects.create(
            cotacao=cotacao,
            enviado_por=request_user,
            destinatario=cotacao.email,
            assunto=assunto,
            corpo=corpo_html,
            enviado_com_sucesso=False
        )

        msg = EmailMultiAlternatives(
            subject=assunto,
            body=corpo_texto,
            from_email=config_email.email,
            to=[cotacao.email],
            reply_to=[config_email.email_resposta or config_email.email]
        )
        msg.attach_alternative(corpo_html, "text/html")

        # O resto da função de envio SMTP continua igual...
        try:
            server = None
            if config_email.porta_smtp == 465:
                context = ssl.create_default_context()
                server = smtplib.SMTP_SSL(config_email.servidor_smtp, config_email.porta_smtp, context=context)
                server.login(config_email.email, config_email.senha)
            elif config_email.usar_tls_smtp:
                server = smtplib.SMTP(config_email.servidor_smtp, config_email.porta_smtp)
                server.starttls()
                server.login(config_email.email, config_email.senha)
            else:
                 server = smtplib.SMTP(config_email.servidor_smtp, config_email.porta_smtp)
                 server.login(config_email.email, config_email.senha)

            server.sendmail(msg.from_email, msg.recipients(), msg.message().as_bytes())
            server.quit()

            salvar_email_enviado(config_email, msg)
            email_registro.enviado_com_sucesso = True
            email_registro.save()
            return True
        except Exception as e:
            email_registro.erro = str(e)
            email_registro.save()
            raise Exception(f'Erro ao enviar email: {e}')

    except Exception as e:
        raise Exception(f'Erro ao preparar email: {str(e)}')


def gerar_mensagem_whatsapp(cotacao, empresa_usuario, user, numero_proposta):
    """
    (VERSÃO UNIFICADA COM MÚLTIPLOS ESTILOS)
    Gera a mensagem de PROPOSTA formatada para WhatsApp.
    A função agora decide internamente qual formato usar.
    """
    # === ESCOLHA O ESTILO DA MENSAGEM AQUI ===
    # Formato 1: Original
    # Formato 2: Visual com emojis (o que fizemos antes)
    # Formato 3: Compacto e Direto
    # Formato 4: Blocos de Informação
    # Formato 5: Minimalista e Elegante
    
    # Defina aqui os IDs e o formato desejado para cada empresa
    if empresa_usuario and empresa_usuario.id in [1, 2, 4]:
        formato_escolhido = 4
    else:
        formato_escolhido = 1 # Formato padrão para as outras empresas
        
    # ==========================================

    # Dados comuns para todos os formatos
    contato_nome = cotacao.contato or 'Cliente'
    valor_proposta = f"R$ {float(cotacao.valor_proposta):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if cotacao.valor_proposta else 'R$ 0,00'
    hora_atual = datetime.now().hour
    if 12 <= hora_atual < 18:
        saudacao_periodo = random.choice(SPINTAX_VARIATIONS["saudacao_hora"]["tarde"])
    elif hora_atual >= 18:
        saudacao_periodo = random.choice(SPINTAX_VARIATIONS["saudacao_hora"]["noite"])
    else:
        saudacao_periodo = random.choice(SPINTAX_VARIATIONS["saudacao_hora"]["manha"])

    assinatura = obter_assinatura_digital(user, empresa_usuario)
    prazo_coleta_original = cotacao.prazo_coleta or 'A COMBINAR'
    prazo_coleta_formatado = prazo_coleta_original[3:] if prazo_coleta_original.upper().startswith('DE ') else prazo_coleta_original

    IDS_EMPRESA_SEM_CARTAO = [5, 9, 10, 11, 12, 13, 15, 16, 17, 18, 20]
    if empresa_usuario.id == 19:
        texto_pagamento = "PIX, TRANSFERENCIA BANCÁRIA E BOLETO MEDIANTE ANÁLISE"
    elif empresa_usuario.id == 8:
        texto_pagamento = "PIX E CARTÃO"
    elif empresa_usuario.id in IDS_EMPRESA_SEM_CARTAO:
        texto_pagamento = "PIX E BOLETO"
    else:
        texto_pagamento = "PIX, BOLETO E CARTÃO"
    
    # --- CONSTRÓI A MENSAGEM COM BASE NO FORMATO ESCOLHIDO ---

    # FORMATO 2: VISUAL (O que fizemos antes, com muito espaço)
    if formato_escolhido == 2:
        mensagem = f"""👋 Olá, *{contato_nome}*! {saudacao_periodo}!

{random.choice(SPINTAX_VARIATIONS['proposta_v2']['apresentacao']).format(empresa=empresa_usuario.nome_full)}

📝 *Resumo da Proposta Nº {numero_proposta}*
 
 🚚 *Rota*
 De: *{cotacao.origem or 'NÃO INFORMADO'}*
 Para: *{cotacao.destino or 'NÃO INFORMADO'}*
 
 💰 *Valor do Frete*
 *{valor_proposta}*
 
 🗓️ *Prazos*
 Coleta: *{prazo_coleta_formatado}*
 Entrega: *{cotacao.prazo_entrega or 'A COMBINAR'}*
 
 💳 *Pagamento*
 Formas: *{texto_pagamento}*
"""
        if cotacao.tipo_frete and cotacao.tipo_frete.strip().lower() != 'mudanças':
            mensagem += "\n\n🛡️ *Seguro Incluso*: _Cobertura para roubo, acidente e avarias._"
        if empresa_usuario.id in [2, 4, 16, 17]:
            mensagem += "\n\n👷‍♂️ _Proposta livre de carga e descarga (serviços por conta do cliente)._"
        mensagem += f"""

Estamos à disposição para qualquer ajuste!
*{random.choice(SPINTAX_VARIATIONS['proposta_v2']['chamada_negociacao'])}* 😉

{assinatura}
"""
        return mensagem

    # FORMATO 3: COMPACTO E DIRETO
    elif formato_escolhido == 3:
        mensagem = f"👋 Olá, *{contato_nome}*! {saudacao_periodo}!\n"
        mensagem += f"{random.choice(SPINTAX_VARIATIONS['proposta_v2']['apresentacao']).format(empresa=empresa_usuario.nome_full)}\n\n"
        mensagem += f"📝 *Proposta Nº {numero_proposta}*\n"
        mensagem += f"• *Rota:* {cotacao.origem or 'N/A'} para {cotacao.destino or 'N/A'}\n"
        mensagem += f"• *Valor:* {valor_proposta}\n"
        mensagem += f"• *Coleta:* {prazo_coleta_formatado}\n"
        mensagem += f"• *Entrega:* {cotacao.prazo_entrega or 'A COMBINAR'}\n"
        mensagem += f"• *Pagamento:* {texto_pagamento}\n"
        if cotacao.tipo_frete and cotacao.tipo_frete.strip().lower() != 'mudanças':
            mensagem += "• *Seguro:* Cobertura para roubo, acidente e avarias.\n"
        if empresa_usuario.id in [2, 4, 16, 17]:
            mensagem += "\n👷‍♂️ _Serviços de carga e descarga por conta do cliente._"
        mensagem += f"\n\nPrecisa de algum ajuste? Me avise! 😉\n\n{assinatura}"
        return mensagem

    # FORMATO 4: BLOCOS DE INFORMAÇÃO
    elif formato_escolhido == 4:
        mensagem = f"👋 Olá, *{contato_nome}*! {saudacao_periodo}!\n"
        mensagem += f"{random.choice(SPINTAX_VARIATIONS['proposta']['introducao']).format(empresa=empresa_usuario.nome_full)}\n\n"
        mensagem += f"--- *PROPOSTA Nº {numero_proposta}* ---\n\n"
        mensagem += f"🚚 *ROTA*\nDe: *{cotacao.origem or 'N/A'}*\nPara: *{cotacao.destino or 'N/A'}*\n\n"
        mensagem += f"💰 *VALOR E PRAZOS*\nFrete: *{valor_proposta}*\nColeta: *{prazo_coleta_formatado}*\nEntrega: *{cotacao.prazo_entrega or 'A COMBINAR'}*\n\n"
        mensagem += f"💳 *PAGAMENTO E SERVIÇOS*\nFormas: *{texto_pagamento}*\n"
        if cotacao.tipo_frete and cotacao.tipo_frete.strip().lower() != 'mudanças':
            mensagem += "Seguro: *ROUBO, ACIDENTE E AVARIAS*\n"
        if empresa_usuario.id in [2, 4, 16, 17]:
            mensagem += "_Serviços de carga e descarga por conta do cliente._\n"
        mensagem += "---------------------------------\n\n"
        mensagem += f"Se os valores ou prazos não baterem, me chame para negociarmos. 😉\n\n{assinatura}"
        return mensagem

    # FORMATO 5: MINIMALISTA E ELEGANTE
    elif formato_escolhido == 5:
        mensagem = f"Olá, {contato_nome}, tudo bem? {saudacao_periodo}.\n\n"
        mensagem += f"Segue sua proposta *Nº {numero_proposta}* com a *{empresa_usuario.nome_full}* para o transporte de *{cotacao.origem}* para *{cotacao.destino}*.\n\n"
        mensagem += f"O valor do frete é de *{valor_proposta}*, com prazo de coleta em *{prazo_coleta_formatado}* e entrega em *{cotacao.prazo_entrega or 'A COMBINAR'}*. "
        mensagem += f"As formas de pagamento são *{texto_pagamento}*.\n\n"
        if cotacao.tipo_frete and cotacao.tipo_frete.strip().lower() != 'mudanças':
            mensagem += "Esta proposta inclui seguro contra roubo, acidente e avarias. "
        if empresa_usuario.id in [2, 4, 16, 17]:
            mensagem += "_O valor não inclui ajudantes para carga/descarga._\n\n"
        mensagem += f"{random.choice(SPINTAX_VARIATIONS['proposta_v2']['chamada_negociacao'])}\n\nAtenciosamente,\n{assinatura}"
        return mensagem

    # FORMATO 1: ORIGINAL (Fallback)
    else:
        partes_da_mensagem = [
            f"{random.choice(SPINTAX_VARIATIONS['proposta']['saudacao_contato']).format(contato=contato_nome)} {saudacao_periodo}!",
            "",
            random.choice(SPINTAX_VARIATIONS['proposta']["introducao"]).format(empresa=empresa_usuario.nome_full),
            "",
            f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_proposta'])} *{numero_proposta}*",
            f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_origem'])} *{cotacao.origem or 'NÃO INFORMADO'}*",
            f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_destino'])} *{cotacao.destino or 'NÃO INFORMADO'}*",
            f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_valor'])} *{valor_proposta}*",
            f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_coleta'])} *{prazo_coleta_formatado}*",
            f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_entrega'])} *{cotacao.prazo_entrega or 'A COMBINAR'}*",
        ]
        if cotacao.tipo_frete and cotacao.tipo_frete.strip().lower() != 'mudanças':
            partes_da_mensagem.append(f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_seguro'])} *ROUBO, ACIDENTE E AVARIAS*")
        partes_da_mensagem.append(f"{random.choice(SPINTAX_VARIATIONS['proposta']['detalhes_pagamento'])} *{texto_pagamento}*")
        if empresa_usuario.id in [2, 4, 16, 17]:
            variacoes_carga_descarga = [
                "👷‍♂️ _Valor não inclui ajudantes para carga/descarga._",
                "👷‍♂️ _Serviços de carga e descarga por conta do cliente._",
                "👷‍♂️ _Proposta livre de carga e descarga._"
            ]
            partes_da_mensagem.append(random.choice(variacoes_carga_descarga))
        partes_da_mensagem.extend([
            "",
            f"🫡 {random.choice(SPINTAX_VARIATIONS['proposta']['negociacao'])}",
            "",
            assinatura
        ])
        return "\n".join(partes_da_mensagem)

def gerar_mensagem_whatsapp_v2(cotacao, empresa_usuario, user, numero_proposta):
    """
    Gera um formato alternativo de mensagem de proposta para WhatsApp,
    com estilo mais visual e direto.
    """
    contato_nome = cotacao.contato or 'Cliente'
    valor_proposta = f"R$ {float(cotacao.valor_proposta):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if cotacao.valor_proposta else 'R$ 0,00'
    
    hora_atual = datetime.now().hour
    if 12 <= hora_atual < 18:
        saudacao_periodo = random.choice(SPINTAX_VARIATIONS["saudacao_hora"]["tarde"])
    elif hora_atual >= 18:
        saudacao_periodo = random.choice(SPINTAX_VARIATIONS["saudacao_hora"]["noite"])
    else:
        saudacao_periodo = random.choice(SPINTAX_VARIATIONS["saudacao_hora"]["manha"])

    assinatura = obter_assinatura_digital(user, empresa_usuario)
    prazo_coleta_original = cotacao.prazo_coleta or 'A COMBINAR'
    prazo_coleta_formatado = prazo_coleta_original[3:] if prazo_coleta_original.upper().startswith('DE ') else prazo_coleta_original

    # Lógica de pagamento (mantida da função original)
    IDS_EMPRESA_SEM_CARTAO = [5, 9, 10, 11, 12, 13, 15, 16, 17, 18]
    if empresa_usuario.id == 19:
        texto_pagamento = "PIX, TRANSFERENCIA BANCÁRIA E BOLETO MEDIANTE ANÁLISE"
    elif empresa_usuario.id == 8:
        texto_pagamento = "PIX E CARTÃO"
    elif empresa_usuario.id in IDS_EMPRESA_SEM_CARTAO:
        texto_pagamento = "PIX E BOLETO"
    else:
        texto_pagamento = "PIX, BOLETO E CARTÃO"

    # --- Construção do novo formato da mensagem ---
    mensagem = f"""👋 Olá, *{contato_nome}*! {saudacao_periodo}!

{random.choice(SPINTAX_VARIATIONS['proposta_v2']['apresentacao']).format(empresa=empresa_usuario.nome_full)}

📝 *Resumo da Proposta Nº {numero_proposta}*
 
 🚚 *Rota*
 De: *{cotacao.origem or 'NÃO INFORMADO'}*
 Para: *{cotacao.destino or 'NÃO INFORMADO'}*
 
 💰 *Valor do Frete*
 *{valor_proposta}*
 
 🗓️ *Prazos*
 Coleta: *{prazo_coleta_formatado}*
 Entrega: *{cotacao.prazo_entrega or 'A COMBINAR'}*
 
 💳 *Pagamento*
 Formas: *{texto_pagamento}*
"""

    # Adiciona a linha de seguro se não for mudança
    if cotacao.tipo_frete and cotacao.tipo_frete.strip().lower() != 'mudanças':
        mensagem += "\n\n🛡️ *Seguro Incluso*: _Cobertura para roubo, acidente e avarias._"

    # Adiciona a mensagem sobre carga e descarga para empresas específicas
    if empresa_usuario.id in [2, 4, 16, 17]:
        mensagem += "\n\n👷‍♂️ _Proposta livre de carga e descarga (serviços por conta do cliente)._"

    # Finalização da mensagem
    mensagem += f"""

Estamos à disposição para qualquer ajuste!
*{random.choice(SPINTAX_VARIATIONS['proposta_v2']['chamada_negociacao'])}* 😉

{assinatura}
"""
    return mensagem

def formatar_numero_whatsapp(numero):
    """Formata um número de telefone para o padrão do WhatsApp (55DDDNUMERO)"""
    numero_limpo = re.sub(r'[^\d]', '', numero)
    
    if len(numero_limpo) < 10:
        raise ValueError('Número de telefone muito curto')
    
    if numero_limpo.startswith('550'):
        numero_limpo = '55' + numero_limpo[3:]
    elif numero_limpo.startswith('55'):
        pass
    elif numero_limpo.startswith('0'):
        numero_limpo = '55' + numero_limpo[1:]
    else:
        numero_limpo = '55' + numero_limpo
    
    return numero_limpo

def limpar_texto_unicode(texto):
    try:
        return unicodedata.normalize('NFKC', texto)
    except:
        return texto

def limpar_unicode_para_url(texto):
    if not isinstance(texto, str):
        texto = str(texto)
    texto = unicodedata.normalize('NFKC', texto)
    texto = texto.encode('utf-8', 'ignore').decode('utf-8')
    return texto

def simular_digitando(instance_key, token, phone, proxy_url=None):
    """(VERSÃO COM PROXY) Envia um sinal de 'digitando' para a MegaAPI."""
    url = f"https://apistart03.megaapi.com.br/rest/chat/{instance_key}/presenceUpdateChat"
    numero_com_sufixo = f"{phone}@s.whatsapp.net"
    payload = {"to": numero_com_sufixo, "presence": "composing"}
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {token}"}
    
    proxies = None
    if proxy_url:
        proxies = {"http": proxy_url, "https": proxy_url}

    try:
        session = _get_secure_requests_session()
        session.post(url, json=payload, headers=headers, proxies=proxies, timeout=3)
    except RequestException:
        pass
def _humanizar_envio(instance_key, token, phone, proxy_url=None):
    """Função central de humanização. Simula 'digitando...' por um período visível."""
    log_proxy_status = f"SIM (Endereço: {proxy_url})" if proxy_url else "NÃO"
    logger.info(f"Humanizando envio para {phone}. Usando Proxy: {log_proxy_status}")
    
    tempo_total_digitando = random.uniform(4, 8)
    intervalo_pulsos = 2.5
    
    start_time = time.time()
    while time.time() - start_time < tempo_total_digitando:
        simular_digitando(instance_key, token, phone, proxy_url)
        tempo_restante = tempo_total_digitando - (time.time() - start_time)
        time.sleep(min(intervalo_pulsos, max(0, tempo_restante)))
    
    logger.info(f"'Digitando' simulado por {time.time() - start_time:.1f} segundos.")

def _enviar_via_megaapi_com_retry(url, payload, headers, proxies, cotacao_telefone):
    """
    (VERSÃO INTELIGENTE) Tenta reenviar a requisição APENAS em caso de falhas de rede.
    Falha imediatamente para erros como 'Número Inválido' ou 'Instância Desconectada'.
    """
    TENTATIVAS_MAXIMAS = 3
    for attempt in range(TENTATIVAS_MAXIMAS):
        try:
            session = _get_secure_requests_session()
            response = session.post(url, json=payload, headers=headers, proxies=proxies, timeout=60)
            
            if 400 <= response.status_code < 500:
                try:
                    error_data = response.json()
                    error_message = error_data.get('message', '').lower()
                    
                    if 'instance not logged in' in error_message:
                        logger.error(f"Erro 403 (Instância Desconectada). Abortando. Resposta: {response.text}")
                        raise WhatsAppInstanceDisconnectedException("A conexão com o WhatsApp está offline. Por favor, reconecte.")
                    
                    if 'number not registered' in error_message:
                        logger.error(f"Erro 403 (Número Inválido) para {cotacao_telefone}. Resposta: {response.text}")
                        raise WhatsAppNumberInvalidException(f"O número de WhatsApp '{cotacao_telefone}' parece ser inválido ou não existe.")
                except (ValueError, AttributeError):
                    pass
            
            response.raise_for_status()
            return response.json()

        except (ProxyError, SSLError, RequestException) as e:
            if isinstance(e, (WhatsAppNumberInvalidException, WhatsAppInstanceDisconnectedException)):
                raise e

            logger.warning(f"Tentativa {attempt + 1}/{TENTATIVAS_MAXIMAS} de envio falhou. Erro de rede: {e}")
            if attempt < TENTATIVAS_MAXIMAS - 1:
                pausa_entre_tentativas = random.uniform(3, 7)
                logger.info(f"Aguardando {pausa_entre_tentativas:.1f} segundos antes da próxima tentativa...")
                time.sleep(pausa_entre_tentativas) 
            else:
                logger.error(f"Todas as {TENTATIVAS_MAXIMAS} tentativas de envio falharam devido a erros de rede.")
                raise

def enviar_whatsapp_megaapi(instance_key, token, phone, message, proxy_url=None, cotacao=None):
    """Envia MENSAGEM DE TEXTO pela MegaAPI usando a lógica de retry inteligente."""
    _humanizar_envio(instance_key, token, phone, proxy_url)
    
    url = f"https://apistart03.megaapi.com.br/rest/sendMessage/{instance_key}/text"
    payload = {"messageData": {"to": phone, "text": message, "linkPreview": False}}
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {token}"}
    proxies = {"http": proxy_url, "https": proxy_url} if proxy_url else None

    try:
        return _enviar_via_megaapi_com_retry(url, payload, headers, proxies, cotacao.telefone)
    except Exception as e:
        logger.error(f"Erro ao enviar WhatsApp (texto) via MegaAPI: {e}")
        raise


def enviar_whatsapp(request, cotacao, mensagem=None, **kwargs):
    user = request.user
    perfil = getattr(user, 'perfil', None)
    if not perfil: raise Exception("Usuário sem perfil.")
    if not cotacao.telefone: raise Exception('Número de WhatsApp não cadastrado')
    
    telefone_formatado = formatar_numero_whatsapp(cotacao.telefone)

    if not mensagem:
        mensagem = gerar_mensagem_whatsapp(cotacao, perfil.empresa, user, gerar_numero_proposta(cotacao))

    if perfil.tem_api_whatsapp():
        if perfil.api_provider == 'MEGAAPI':
            # === DESPACHANTE MEGAAPI ===
            return enviar_whatsapp_megaapi(
                instance_key=perfil.api_credentials.get('instance_key'),
                token=perfil.api_credentials.get('token'),
                phone=telefone_formatado, message=mensagem,
                proxy_url=perfil.proxy_url, cotacao=cotacao
            )
        elif perfil.api_provider == 'SELF_HOSTED':
            # === DESPACHANTE API PRÓPRIA (SELF_HOSTED) ===
            # Passa 'cotacao' e 'user' para a próxima função
            return enviar_texto_self_hosted(
                credentials=perfil.api_credentials,
                phone=telefone_formatado,
                message=mensagem,
                cotacao=cotacao, 
                user=user
            )
        else:
            raise Exception(f"Provedor de API '{perfil.api_provider}' não suportado.")
    else:
        whatsapp_url = f"https://web.whatsapp.com/send?phone={telefone_formatado}&text={quote(mensagem)}"
        return {'redirect': whatsapp_url}
def salvar_email_enviado(config_email, msg):
    import imaplib
    import time

    raw_msg = msg.message()
    eml_bytes = raw_msg.as_bytes()

    imap = imaplib.IMAP4_SSL(config_email.servidor_imap, config_email.porta_imap or 993)
    imap.login(config_email.email, config_email.senha)

    sent_folder = 'INBOX.Sent'
    imap.append(sent_folder, '(\Seen)', imaplib.Time2Internaldate(time.time()), eml_bytes)
    imap.logout()

LIBREOFFICE_PATH = '/usr/bin/libreoffice'

def converter_docx_para_pdf(input_path, output_dir=None):
    """
    Converte um arquivo .docx para .pdf usando LibreOffice,
    sendo resiliente a códigos de saída de aviso se o arquivo for criado.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Arquivo de entrada não encontrado: {input_path}")

    if not output_dir:
        output_dir = os.path.dirname(input_path)

    # Determina o nome do arquivo de saída esperado
    output_filename = os.path.splitext(os.path.basename(input_path))[0] + '.pdf'
    expected_output_path = os.path.join(output_dir, output_filename)

    command = [
        '/usr/bin/libreoffice',
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', output_dir,
        input_path
    ]

    # Define o ambiente para garantir que o LibreOffice encontre seus binários
    env = os.environ.copy()
    env["HOME"] = os.path.expanduser("~") # Adicionar o diretório home pode resolver problemas de permissão
    env["PATH"] = "/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin"

    result = subprocess.run(command, capture_output=True, text=True, env=env)

    # Lógica de verificação aprimorada
    if result.returncode != 0:
        # O comando retornou um erro. Vamos verificar se o PDF foi criado mesmo assim.
        if not os.path.exists(expected_output_path):
            # ERRO REAL: O comando falhou E o arquivo não foi criado.
            # Lançamos uma exceção com detalhes para depuração.
            error_details = (
                f"LibreOffice falhou com o código de saída {result.returncode} E o arquivo PDF não foi criado.\n"
                f"Comando: {' '.join(command)}\n"
                f"Stderr: {result.stderr}\n"
                f"Stdout: {result.stdout}"
            )
            logger.error(error_details)
            raise subprocess.CalledProcessError(
                result.returncode, result.args, output=result.stdout, stderr=result.stderr
            )
        else:
            # AVISO: O comando retornou um erro, mas o PDF foi criado.
            # Apenas registramos o aviso e continuamos.
            warning_details = (
                f"LibreOffice retornou o código de saída {result.returncode}, mas o PDF foi criado com sucesso em '{expected_output_path}'. "
                f"Isso geralmente indica um aviso.\n"
                f"Stderr: {result.stderr}"
            )
            logger.warning(warning_details)

    # Se chegamos aqui, ou o comando teve sucesso (código 0) ou teve um aviso mas gerou o arquivo.
    # Em ambos os casos, a operação foi bem-sucedida para o nosso propósito.
    return expected_output_path

def enviar_whatsapp_com_pdf(request, cotacao, caminho_pdf, nome_arquivo_pdf, legenda=None):
    """
    (VERSÃO COM PROXY) Função SÍNCRONA que envia WhatsApp com PDF.
    Agora delega para a mesma lógica das tasks, garantindo a humanização.
    """
    user = request.user
    # A função `enviar_whatsapp_com_pdf_task` já contém toda a lógica
    # de humanização, proxy e tratamento de erro. Vamos reutilizá-la.
    return enviar_whatsapp_com_pdf_task(user, cotacao, caminho_pdf, nome_arquivo_pdf, legenda)
    
def preparar_url_whatsapp(numero, mensagem):
    """Prepara URL para abrir conversa no WhatsApp com mensagem pré-definida"""
    numero = ''.join(filter(str.isdigit, numero))
    
    if not numero.startswith('55') and len(numero) <= 11:
        numero = '55' + numero
    
    return f"https://wa.me/{numero}?text={quote(mensagem)}"

def gerar_mensagem_solicitacao_medidas(cotacao, user=None):
    """ (VERSÃO CORRIGIDA - COM ORIGEM/DESTINO E AJUSTES)
    Gera a mensagem de SOLICITAÇÃO DE MEDIDAS para o WhatsApp com variações e instruções claras.
    """
    contato = sanitize_text(cotacao.contato or 'Cliente')
    numero_proposta = gerar_numero_proposta(cotacao)
    origem = sanitize_text(cotacao.origem or 'NÃO INFORMADO')
    destino = sanitize_text(cotacao.destino or 'NÃO INFORMADO')
    
    # Monta a introdução com variações
    introducao = random.choice(SPINTAX_VARIATIONS['solicitacao_medidas']['saudacao']).format(contato=contato)
    pedido = random.choice(SPINTAX_VARIATIONS['solicitacao_medidas']['pedido']).format(numero_proposta=numero_proposta)
    
    # Cria uma linha padrão com a rota para ser usada em todas as mensagens
    info_rota = f"🚚 Rota: *{origem}* para *{destino}*"

    # Dados da cotação
    volumes = f"*{sanitize_text(cotacao.volumes) or 'NÃO INFORMADO'}*"
    modelo = f"*{'NÃO INFORMADO'}*"
    ano = f"*{'NÃO INFORMADO'}*"
    valor_mercadoria = cotacao.valor_mercadoria or 0
    cubagem_valor = cotacao.cubagem or 0
    valor = f"*R$ {float(valor_mercadoria):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') + "*"
    cubagem = f"*{cubagem_valor:.2f} m³".replace('.', ',') + "*"
    observacao = sanitize_text(cotacao.observacao or 'NENHUMA')

    empresa_para_assinatura = cotacao.empresa
    if user and hasattr(user, 'perfil') and user.perfil and user.perfil.empresa:
        empresa_para_assinatura = user.perfil.empresa
    assinatura = obter_assinatura_digital(user, empresa_para_assinatura)
    
    # Lógica para os diferentes tipos de frete (agora incluindo a info_rota)
    if cotacao.tipo_frete == 'Mudanças':
        # --- INÍCIO DA CORREÇÃO ---
        # A variável {pedido} foi adicionada aqui
        corpo = f"""{introducao}

{pedido}

{info_rota}

Poderia por gentileza enviar a *lista de itens da mudança*? Se não tiver, pode ser fotos dos móveis/aparelhos, ok?

{random.choice(SPINTAX_VARIATIONS['solicitacao_medidas']['despedida'])}

{assinatura}"""
        # --- FIM DA CORREÇÃO ---
    
    elif cotacao.tipo_frete == 'Moto':
        corpo = f"""{introducao}

{pedido}

{info_rota}

Nº VOLUMES: {volumes}
MODELO: {modelo}
ANO: {ano}
VALOR NF: {valor}

OBSERVAÇÃO:
{observacao}

{random.choice(SPINTAX_VARIATIONS['solicitacao_medidas']['despedida'])}

{assinatura}"""

    else: # Frete Geral (Cargas)
        instrucao_clara = random.choice(SPINTAX_VARIATIONS['solicitacao_medidas']['instrucao_geral'])
        
        corpo = f"""{introducao}

{pedido}

{info_rota}

_{instrucao_clara}_

Nº VOLUMES: {volumes}
CUBAGEM: {cubagem}
VALOR NF: {valor}

OBSERVAÇÃO:
{observacao}

{random.choice(SPINTAX_VARIATIONS['solicitacao_medidas']['despedida'])}

{assinatura}"""

    return corpo

# --- NOVA FUNÇÃO ADICIONADA ---
def gerar_html_solicitacao_medidas(cotacao, numero_proposta, empresa, assinatura_digital):
    """
    Renderiza o template HTML para o e-mail de solicitação de medidas.
    """
    context = {
        'cotacao': cotacao,
        'numero_proposta': numero_proposta,
        'empresa': empresa,
        'assinatura_digital': assinatura_digital,
    }
    # Usa o novo template para renderizar o HTML
    return render_to_string('cotacoes/email_solicitacao_medidas.html', context)


def gerar_mensagem_pagamento(cotacao, empresa, user, status_mercadoria, percentual, mensagem_adicional=''):
    """ (NOVA VERSÃO HUMANIZADA)
    Gera a mensagem de SOLICITAÇÃO DE PAGAMENTO com variações.
    """
    dados = {} # Usado para armazenar os dados para o e-mail
    
    numero_proposta = gerar_numero_proposta(cotacao)
    cliente = cotacao.contato or "Cliente"
    
    valor_total = cotacao.valor_proposta or 0
    if percentual == 50:
        valor_parcial = float(valor_total) / 2
        valor_exibido = f"R$ {valor_parcial:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        referencia_percentual = f"(referente a 50%)"
    else:
        valor_exibido = f"R$ {float(valor_total):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        referencia_percentual = f"(referente a 100%)"
    
    razao_social = empresa.razao_social or empresa.nome_full or empresa.nome
    banco = empresa.banco or "Não informado"
    agencia_conta = empresa.agencia_conta or "Não informado"
    pix = empresa.pix or "Não informado"
    assinatura_digital = obter_assinatura_digital(user, empresa)
    
    texto_status = random.choice(SPINTAX_VARIATIONS['solicitacao_pagamento']['status_mercadoria'])

    # Monta a mensagem para o WhatsApp com variações
    mensagem_whatsapp = (
        f"{random.choice(SPINTAX_VARIATIONS['solicitacao_pagamento']['saudacao']).format(cliente=cliente)} "
        f"{random.choice(SPINTAX_VARIATIONS['solicitacao_pagamento']['noticia_boa'])}\n\n"
        f"{texto_status}\n\n"
        f"🏳 Origem....: *{cotacao.origem or 'Não informado'}*\n"
        f"🏁 Destino....: *{cotacao.destino or 'Não informado'}*\n"
        f"📄 Proposta..: *{numero_proposta}*\n\n"
        f"Abaixo seguem os dados para o pagamento:\n\n"
        f"💰 Valor........: *{valor_exibido} {referencia_percentual}*\n"
        f"🏧 PIX............: *{pix}*\n"
        f"💳 Favorec...: *{razao_social}*\n"
        f"🏦 Banco......: *{banco} | Agência/Conta: {agencia_conta}*\n\n"
    )
    
    if mensagem_adicional:
        mensagem_whatsapp += f"Observação: {mensagem_adicional}\n\n"
        
    mensagem_whatsapp += f"{random.choice(SPINTAX_VARIATIONS['solicitacao_pagamento']['pedido_comprovante'])}\n\n"
    mensagem_whatsapp += f"{random.choice(SPINTAX_VARIATIONS['solicitacao_pagamento']['despedida'])}\n\n{assinatura_digital}"

    return {
        "mensagem_whatsapp": mensagem_whatsapp,
        "numero_proposta": numero_proposta,
        "valor_exibido": valor_exibido,
        "referencia_percentual": referencia_percentual,
        "dados_bancarios": {
            "banco": banco,
            "agencia_conta": agencia_conta,
            "pix": pix,
            "razao_social": razao_social
        },
        "texto_status": texto_status,
        "cliente": cliente,
        "mensagem_adicional": mensagem_adicional,
    }

# --- ADICIONE ESTA NOVA FUNÇÃO ---
def gerar_html_solicitacao_pagamento(cotacao, numero_proposta, empresa, assinatura_digital, dados_pagamento):
    """
    Renderiza o template HTML para o e-mail de solicitação de pagamento.
    """
    context = {
        'cotacao': cotacao,
        'numero_proposta': numero_proposta,
        'empresa': empresa,
        'assinatura_digital': assinatura_digital,
        'dados': dados_pagamento,
    }
    return render_to_string('cotacoes/email_solicitacao_pagamento.html', context)

def formatar_whatsapp_para_exibicao(numero):
    """
    Formata um número de telefone do padrão '55 11 9NNNN NNNN' 
    para o formato de exibição '(11) 9NNNN-NNNN'.
    """
    if not numero:
        return ""  # Retorna vazio se não houver número

    try:
        # Remove todos os caracteres não numéricos
        numero_limpo = re.sub(r'[^\d]', '', numero)

        # Se o número começar com 55 e tiver 13 dígitos (55 + DDD + 9 dígitos)
        if numero_limpo.startswith('55') and len(numero_limpo) == 13:
            ddd = numero_limpo[2:4]
            parte1 = numero_limpo[4:9]
            parte2 = numero_limpo[9:]
            return f"({ddd}) {parte1}-{parte2}"
        
        # Se o número tiver 11 dígitos (DDD + 9 dígitos)
        elif len(numero_limpo) == 11:
            ddd = numero_limpo[0:2]
            parte1 = numero_limpo[2:7]
            parte2 = numero_limpo[7:]
            return f"({ddd}) {parte1}-{parte2}"

        # Se não corresponder, retorna o número original de forma segura
        return numero
    except Exception:
        # Em caso de qualquer erro, retorna o número original
        return numero

def enviar_apenas_email_task(cotacao_id, user_id, output_pdf_path=None, output_pdf_filename=None):
    """
    Versão da função de envio de e-mail.
    OTIMIZADO: Aceita um PDF já gerado ou gera um novo se necessário.
    """
    logger.info(f"[TASK] Iniciando envio de e-mail para cotação ID: {cotacao_id}")
    User = get_user_model()
    try:
        user = User.objects.get(id=user_id)
        cotacao = Cotacao.objects.get(id=cotacao_id)
        
        if not hasattr(user, 'perfil') or not user.perfil.empresa:
            raise Exception('Usuário sem perfil ou empresa associada.')

        empresa_usuario = user.perfil.empresa
        
        if not cotacao.email or '@' not in cotacao.email:
            logger.warning(f"[TASK] Cotação {cotacao_id} sem e-mail válido. Pulando envio de e-mail.")
            return

        numero_proposta = gerar_numero_proposta(cotacao)

        # Se o PDF não foi passado como parâmetro, esta função o gera.
        # Caso contrário, ela pula esta etapa demorada.
        if not output_pdf_path or not output_pdf_filename:
            logger.warning(f"[TASK] Gerando PDF de fallback para e-mail da cotação {cotacao_id}.")
            replacements = criar_replacements(cotacao, numero_proposta, user)
            if not empresa_usuario.template_proposta:
                raise Exception(f"Template de proposta não configurado para a empresa ID: {empresa_usuario.id}.")
            template_path = empresa_usuario.template_proposta.path
            output_pdf_path, output_pdf_filename, _ = gerar_proposta_word(
                cotacao=cotacao,
                replacements=replacements,
                empresa_usuario=empresa_usuario,
                template_path=template_path,
                user=user
            )

        config_email = ConfiguracaoEmail.objects.get(empresa=empresa_usuario)
        enviar_email_proposta(
            config_email,
            cotacao,
            output_pdf_path,
            output_pdf_filename,
            numero_proposta,
            empresa_usuario,
            user
        )

        # Atualiza status corretamente
        cotacao.refresh_from_db()
        if 'Enviado Email' not in (cotacao.status_envio or ''):
            if cotacao.status_envio and 'Enviado Whats' in cotacao.status_envio:
                cotacao.status_envio = 'Enviado Whats + Email'
            else:
                cotacao.status_envio = 'Enviado Email'
        
        cotacao.proposta_gerada = os.path.join('propostas', output_pdf_filename)
        cotacao.save()
        logger.info(f"[TASK] E-mail para cotação {cotacao_id} enviado com sucesso. Status atualizado para: {cotacao.status_envio}")

    except (Cotacao.DoesNotExist, User.DoesNotExist, ConfiguracaoEmail.DoesNotExist) as e:
        logger.error(f"[TASK] Erro de objeto não encontrado ao enviar e-mail para cotação {cotacao_id}: {e}", exc_info=True)
    except Exception as e:
        logger.error(f"[TASK] Erro inesperado ao enviar e-mail para cotação {cotacao_id}: {e}", exc_info=True)
        raise


def enviar_apenas_whatsapp_task(cotacao_id, user_id, caminho_pdf=None, nome_arquivo_pdf=None):
    """
    Versão da função de envio de WhatsApp.
    OTIMIZADO: Aceita um PDF já gerado ou gera um novo se necessário.
    """
    logger.info(f"[TASK] Iniciando envio de WhatsApp para cotação ID: {cotacao_id}")
    User = get_user_model()
    try:

        user = User.objects.get(id=user_id)
        cotacao = Cotacao.objects.get(id=cotacao_id)

        if not hasattr(user, 'perfil') or not user.perfil.empresa:
            raise Exception('Usuário sem perfil configurado')
        
        if not (hasattr(user.perfil, 'tem_api_whatsapp') and user.perfil.tem_api_whatsapp()):
            logger.warning(f"[TASK] Usuário {user.username} não tem API de WhatsApp configurada. Pulando envio para cotação {cotacao_id}.")
            return

        empresa_usuario = user.perfil.empresa
        
        if not cotacao.telefone or len(cotacao.telefone) < 8:
            logger.warning(f"[TASK] Cotação {cotacao_id} sem telefone válido. Pulando envio de WhatsApp.")
            return

        numero_proposta = gerar_numero_proposta(cotacao)

        # Se o PDF não foi passado como parâmetro, esta função o gera.
        # Caso contrário, ela pula esta etapa demorada.
        if not caminho_pdf or not nome_arquivo_pdf:
            logger.warning(f"[TASK] Gerando PDF de fallback para WhatsApp da cotação {cotacao_id}.")
            replacements = criar_replacements(cotacao, numero_proposta, user)
            if not empresa_usuario.template_proposta:
                raise Exception(f"Template de proposta não configurado para a empresa ID: {empresa_usuario.id}.")
            template_path = empresa_usuario.template_proposta.path
            caminho_pdf, nome_arquivo_pdf, _ = gerar_proposta_word(
                cotacao=cotacao,
                replacements=replacements,
                empresa_usuario=empresa_usuario,
                template_path=template_path,
                user=user
            )

        legenda = gerar_mensagem_whatsapp(cotacao, empresa_usuario, user, numero_proposta)

        resultado = enviar_whatsapp_com_pdf_task(
            user=user,
            cotacao=cotacao,
            caminho_pdf=caminho_pdf,
            nome_arquivo_pdf=nome_arquivo_pdf,
            legenda=legenda
        )

        # Atualiza status corretamente
        cotacao.refresh_from_db()
        if 'Enviado Whats' not in (cotacao.status_envio or ''):
            if cotacao.status_envio and 'Enviado Email' in cotacao.status_envio:
                cotacao.status_envio = 'Enviado Whats + Email'
            else:
                cotacao.status_envio = 'Enviado Whats'
        
        cotacao.proposta_gerada = os.path.join('propostas', nome_arquivo_pdf)
        cotacao.save()
        logger.info(f"[TASK] WhatsApp para cotação {cotacao_id} enviado com sucesso. Resultado API: {resultado}. Status atualizado para: {cotacao.status_envio}")

    except (Cotacao.DoesNotExist, User.DoesNotExist) as e:
        logger.error(f"[TASK] Erro de objeto não encontrado ao enviar WhatsApp para cotação {cotacao_id}: {e}", exc_info=True)
    except Exception as e:
        logger.error(f"[TASK] Erro inesperado ao enviar WhatsApp para cotação {cotacao_id}: {e}", exc_info=True)
        raise


def enviar_whatsapp_com_pdf_task(user, cotacao, caminho_pdf, nome_arquivo_pdf, legenda=None):
    perfil = getattr(user, 'perfil', None)
    if not perfil: raise Exception("Usuário sem perfil.")
    if not cotacao.telefone: raise Exception('Número de WhatsApp não cadastrado.')
    
    telefone_puro = formatar_numero_whatsapp(cotacao.telefone)
    
    with open(caminho_pdf, "rb") as f:
        pdf_base64 = base64.b64encode(f.read()).decode('utf-8')
    
    if not legenda:
        legenda = gerar_mensagem_whatsapp(cotacao, perfil.empresa, user, gerar_numero_proposta(cotacao))

    if perfil.api_provider == 'MEGAAPI':
        # (Lógica da MegaAPI permanece a mesma)
        _humanizar_envio(
            instance_key=perfil.api_credentials.get('instance_key'), 
            token=perfil.api_credentials.get('token'),          
            phone=telefone_puro,
            proxy_url=perfil.proxy_url
        )
        payload = { "messageData": { "to": f"{telefone_puro}@s.whatsapp.net", "base64": f"data:application/pdf;base64,{pdf_base64}", "fileName": nome_arquivo_pdf, "type": "document", "caption": legenda, "mimeType": "application/pdf" }}
        url = f"https://apistart03.megaapi.com.br/rest/sendMessage/{perfil.api_credentials.get('instance_key')}/mediaBase64"
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {perfil.api_credentials.get('token')}"}
        proxies = {"http": perfil.proxy_url, "https": perfil.proxy_url} if perfil.proxy_url else None
        return _enviar_via_megaapi_com_retry(url, payload, headers, proxies, cotacao.telefone)
        
    elif perfil.api_provider == 'SELF_HOSTED':
        # === DESPACHANTE API PRÓPRIA (SELF_HOSTED) PARA PDF ===
        credentials = perfil.api_credentials
        
        # ⭐️ PASSO 1: VERIFICAR O NÚMERO ⭐️
        try:
            numero_existe = _check_number_self_hosted(credentials, telefone_puro)
        except Exception as e:
            raise e # Propaga o erro (ex: API desconectada)

        if not numero_existe:
            logger.warning(f"Bloqueado envio de PDF para {telefone_puro} (Instância {credentials.get('port')}): Número não existe no WhatsApp.")
            raise WhatsAppNumberInvalidException(f"O número {telefone_puro} não existe no WhatsApp.")

        # ⭐️ PASSO 2: ENVIAR (Se o número existe) ⭐️
        port = credentials.get('port')
        api_url = f"http://127.0.0.1:{port}/send-pdf" 
        
        payload = { 
            "number": telefone_puro, 
            "base64": pdf_base64, 
            "fileName": nome_arquivo_pdf, 
            "caption": legenda 
        }
        headers = {"Content-Type": "application/json"}
        
        try:
            response = requests.post(api_url, json=payload, headers=headers, timeout=40)
            response.raise_for_status()
            
            # ⭐️ LÓGICA DE LOG (Simplificada) ⭐️
            responseData = response.json()
            try:
                EmailEnviado.objects.create(
                    cotacao=cotacao,
                    enviado_por=user,
                    destinatario=telefone_puro,
                    assunto=f"WhatsApp PDF Enviado (SELF_HOSTED)",
                    corpo=legenda or "PDF Enviado",
                    enviado_com_sucesso=True, # Sucesso no envio
                )
            except Exception as e:
                logger.warning(f"Falha ao salvar log do WhatsApp PDF: {e}")
            
            return responseData
            
        except RequestException as e:
            logger.error(f"Erro ao enviar PDF para API SELF_HOSTED na porta {port}: {e}")
            raise

    else:
        raise Exception(f"Provedor de API '{perfil.api_provider}' inválido para envio de PDF.")
def enviar_proposta_sincrono(cotacao_id, user_id):
    """
    (VERSÃO CORRIGIDA COM TOLERÂNCIA A FALHAS E SEM DEPENDÊNCIA DE TASKS)
    Executa o envio síncrono, garantindo que a falha em um canal
    (e-mail) não impeça a tentativa de envio no outro (WhatsApp).
    """
    User = get_user_model()
    cotacao = Cotacao.objects.get(pk=cotacao_id)
    user = User.objects.get(pk=user_id)
    empresa_usuario = user.perfil.empresa

    resultado = {
        'sucesso_geral': False, 'enviado_email': False, 'enviado_whatsapp': False,
        'erros': [], 'status_final': cotacao.status_envio
    }
    output_pdf_path = None
    output_pdf_filename = None

    try:
        # --- Geração de PDF (Etapa Crítica) ---
        numero_proposta = gerar_numero_proposta(cotacao)
        replacements = criar_replacements(cotacao, numero_proposta, user)

        if not empresa_usuario.template_proposta or not os.path.exists(empresa_usuario.template_proposta.path):
            raise Exception("Template de proposta não configurado ou não encontrado.")

        output_pdf_path, output_pdf_filename, _ = gerar_proposta_word(
            cotacao=cotacao, replacements=replacements, empresa_usuario=empresa_usuario,
            template_path=empresa_usuario.template_proposta.path, user=user
        )

        # --- Envio por Email (com tratamento de erro isolado) ---
        if cotacao.email and '@' in cotacao.email:
            try:
                # Chama a função síncrona que está no próprio utils.py
                enviar_apenas_email_task(cotacao_id, user_id, output_pdf_path=output_pdf_path, output_pdf_filename=output_pdf_filename)
                resultado['enviado_email'] = True
            except Exception as e:
                error_msg = f"Falha ao enviar e-mail: {str(e)}"
                logger.error(error_msg, exc_info=True)
                resultado['erros'].append(error_msg)

        # --- Pausa de segurança ---
        time.sleep(1)

        # --- Envio por WhatsApp (com tratamento de erro isolado) ---
        if cotacao.telefone and hasattr(user, 'perfil') and user.perfil.tem_api_whatsapp():
            try:
                # Chama a função síncrona que está no próprio utils.py
                enviar_apenas_whatsapp_task(cotacao_id, user_id, caminho_pdf=output_pdf_path, nome_arquivo_pdf=output_pdf_filename)
                resultado['enviado_whatsapp'] = True
            except (WhatsAppInstanceDisconnectedException, WhatsAppNumberInvalidException) as e:
                error_msg = str(e)
                logger.warning(f"Falha no envio de WhatsApp: {error_msg}")
                resultado['erros'].append(error_msg)
            except Exception as e:
                error_msg = f"Falha desconhecida ao enviar WhatsApp: {str(e)}"
                logger.error(error_msg, exc_info=True)
                resultado['erros'].append(error_msg)

        # --- Atualização Final do Status ---
        cotacao.refresh_from_db()
        status_atual = cotacao.status_envio or ''

        if resultado['enviado_email'] and 'Enviado Email' not in status_atual:
            status_atual = 'Enviado Whats + Email' if 'Enviado Whats' in status_atual else 'Enviado Email'

        if resultado['enviado_whatsapp'] and 'Enviado Whats' not in status_atual:
            status_atual = 'Enviado Whats + Email' if 'Enviado Email' in status_atual else 'Enviado Whats'

        if resultado['enviado_email'] or resultado['enviado_whatsapp']:
            resultado['sucesso_geral'] = True
            cotacao.status_envio = status_atual
            cotacao.proposta_gerada = os.path.join('propostas', output_pdf_filename)
            cotacao.save()

        resultado['status_final'] = cotacao.status_envio
        return resultado

    except Exception as e:
        logger.critical(f"Erro crítico no envio síncrono da cotação {cotacao_id}: {e}", exc_info=True)
        resultado['erros'].append(str(e))
        return resultado

def gerar_mensagem_coleta_whatsapp(cotacao, numero_proposta):
    """ (NOVA VERSÃO HUMANIZADA)
    Gera a mensagem com o formulário de SOLICITAÇÃO DE COLETA para o WhatsApp com variações.
    """
    introducao = random.choice(SPINTAX_VARIATIONS['solicitacao_coleta']['introducao']).format(numero_proposta=numero_proposta)
    pedido = random.choice(SPINTAX_VARIATIONS['solicitacao_coleta']['pedido'])
    
    # Mantém a lógica específica por empresa
    if cotacao.empresa and cotacao.empresa.id in [9, 10, 18]:
        formulario = f"""*DADOS DO REMETENTE*\n*Nome/Razão Social:* *CPF/CNPJ:*\n*Endereço completo c/ cep:* *Responsável p/ recebimento:*\n*Telefone celular:* *Horários permitidos:* 00:00 às 00:00\n*Coleta em fds/feriado:* ( ) sim ( ) não\n*Ajudante Carga:* ( ) sim ( ) não, quantos?\n\n📦 *DADOS DO DESTINATÁRIO*\n*Nome/Razão Social:*\n*CPF/CNPJ:*\n*Endereço completo c/ cep:* *Responsável p/ recebimento:*\n*Telefone celular:*\n*Horários permitidos:* 00:00 às 00:00\n*Entrega em fds/feriado:* ( ) sim ( ) não\n*Ajudante Descarga:* ( ) sim ( ) não, quantos?\n\n*Dados pagador do frete:*\n*CNPJ/CPF:*"""
    elif cotacao.empresa and cotacao.empresa.id == 13:
        formulario = f"""Data que o material estará liberado para coletar: \nHorário de funcionamento do local de coleta (informar horário de almoço):\nHorário de funcionamento do local de entrega (informar horário de almoço):\nCPF ou CNPJ do remetente:\nCPF ou CNPJ do destinatario: \nCPF ou CNPJ do pagador do frete: \nEndereço completo de coleta: \nEndereço completo de entrega: \nNome e Telefone para contato do remetente: \nNome e Telefone para contato do destinatário:\nTipo de material: \nPreferência de veículo aberto ou fechado."""
    else: # Formulário Padrão
        formulario = f"""📦 *DADOS DO REMETENTE*\n*Nome/Razão Social:* *CPF/CNPJ:*\n*Endereço completo c/ cep:* *Responsável p/ recebimento:*\n*Telefone celular:* *Horários permitidos:* 00:00 às 00:00\n*Coleta em fds/feriado:* ( ) sim ( ) não\n*Ajudante Carga:* ( ) sim ( ) não, quantos?\n\n📦 *DADOS DO DESTINATÁRIO*\n*Nome/Razão Social:*\n*CPF/CNPJ:*\n*Endereço completo c/ cep:* *Responsável p/ recebimento:*\n*Telefone celular:*\n*Horários permitidos:* 00:00 às 00:00\n*Entrega em fds/feriado:* ( ) sim ( ) não\n*Ajudante Descarga:* ( ) sim ( ) não, quantos?"""
        
    return f"{introducao}\n{pedido}\n\n{formulario}".strip()


def gerar_html_formulario_coleta(cotacao, numero_proposta, empresa, assinatura_digital):
    """
    Renderiza o template HTML para o e-mail de solicitação de coleta,
    garantindo um visual profissional e consistente.
    """
    context = {
        'cotacao': cotacao,
        'numero_proposta': numero_proposta,
        'empresa': empresa,
        'assinatura_digital': assinatura_digital,
    }

    # --- INÍCIO DA ALTERAÇÃO ---
    # Define qual template usar com base no ID da empresa
    if empresa and empresa.id == 13:
        template_name = 'cotacoes/email_solicitacao_coleta_empresa13_fundamental.html'
    else:
        template_name = 'cotacoes/email_solicitacao_coleta.html'
    # --- FIM DA ALTERAÇÃO ---

    # Usa o template definido para renderizar o HTML
    return render_to_string(template_name, context)


def enviar_texto_self_hosted(credentials, phone, message, cotacao=None, user=None):
    """Envia uma mensagem de texto para a sua própria infraestrutura."""
    
    # ⭐️ PASSO 1: VERIFICAR O NÚMERO ⭐️
    try:
        numero_existe = _check_number_self_hosted(credentials, phone)
    except Exception as e:
        # Se a *verificação* falhar (ex: API down), levanta a exceção
        raise e 

    if not numero_existe:
        # Se o número NÃO existe, levanta a exceção ANTES de enviar
        logger.warning(f"Bloqueado envio de Texto para {phone} (Instância {credentials.get('port')}): Número não existe no WhatsApp.")
        raise WhatsAppNumberInvalidException(f"O número {phone} não existe no WhatsApp.")
    
    # ⭐️ PASSO 2: ENVIAR (Se o número existe) ⭐️
    port = credentials.get('port')
    api_url = f"http://127.0.0.1:{port}/send-text"
    payload = {"number": phone, "message": message}
    headers = {"Content-Type": "application/json"}
    
    proxies = None # Proxy é tratado no api.js

    try:
        response = requests.post(api_url, json=payload, headers=headers, proxies=proxies, timeout=40)
        response.raise_for_status() 
        
        # ⭐️ LÓGICA DE LOG (Simplificada) ⭐️
        responseData = response.json()
        
        if cotacao and user:
            try:
                EmailEnviado.objects.create(
                    cotacao=cotacao,
                    enviado_por=user,
                    destinatario=phone,
                    assunto=f"WhatsApp Texto Enviado (SELF_HOSTED)",
                    corpo=message,
                    enviado_com_sucesso=True, # Sucesso no envio
                )
            except Exception as e:
                logger.warning(f"Falha ao salvar log do WhatsApp (Texto): {e}")

        return responseData
        
    except RequestException as e:
        logger.error(f"Erro ao comunicar com a API interna na porta {port}: {e}")
        
        if e.response is not None and e.response.status_code == 400:
            raise WhatsAppInstanceDisconnectedException("A API está offline ou desconectada. Verifique o Whatsapp.")
        
        raise

def get_whatsapp_api_status(perfil):
    """
    Verifica o status da API Node.js (SELF_HOSTED) para o perfil fornecido.
    Retorna um dicionário com o status.
    """
    provider = perfil.api_provider
    credentials = perfil.api_credentials
    
    if provider == 'SELF_HOSTED':
        port = credentials.get('port')
        instance_name = credentials.get('instance_name')
        
        if not port or not instance_name:
            return {"status": "error", "message": "Credenciais da API SELF_HOSTED incompletas."}

        # Endpoint /status da API Node.js
        api_url = f"http://127.0.0.1:{port}/status"
        
        try:
            # Tenta um GET simples para verificar se a API está de pé
            response = requests.get(api_url, timeout=5)
            response.raise_for_status() # Lança HTTPError para status 4xx/5xx
            
            # A API Node.js retorna o status real (e.g., {"status": "connected"})
            return response.json() 
            
        except RequestException:
            # Erro de conexão ou HTTPError (400, 404, 500)
            return {"status": "disconnected", "message": f"Instância '{instance_name}' (Porta {port}) está offline ou não conectada ao WhatsApp."}

    return {"status": "error", "message": f"Provedor de API '{provider}' não suportado nesta verificação."}    

def _check_number_self_hosted(credentials, phone):
    """
    Chama o endpoint /check-number na API Node.js para verificar se um número
    existe no WhatsApp ANTES de tentar o envio.
    """
    port = credentials.get('port')
    api_url = f"http://127.0.0.1:{port}/check-number"
    payload = {"number": phone}
    headers = {"Content-Type": "application/json"}
    
    try:
        response = requests.post(api_url, json=payload, headers=headers, timeout=10) # Timeout de 10s
        response.raise_for_status() # Lança erro para 4xx/5xx
        
        data = response.json()
        
        if data.get('success'):
            logger.info(f"[Check Number] Verificação para {phone} (Porta {port}): Existe = {data.get('exists')}")
            return data.get('exists') # Retorna True ou False
        else:
            # A API retornou success: false (ex: API não conectada)
            raise WhatsAppInstanceDisconnectedException(data.get('error', 'API de verificação falhou'))
            
    except requests.exceptions.HTTPError as e:
        # Se a API retornar 400 (ex: "Instância desconectada")
        if e.response.status_code == 400:
            try:
                error_data = e.response.json()
                raise WhatsAppInstanceDisconnectedException(error_data.get('error', 'API desconectada (400)'))
            except ValueError:
                pass # Cai no erro genérico abaixo
        
        # Erro 500 ou outro erro HTTP
        logger.error(f"Erro HTTP ao VERIFICAR número na porta {port}: {e}")
        raise Exception(f"Erro HTTP {e.response.status_code} ao verificar número.")
        
    except requests.exceptions.RequestException as e:
        # Erro de conexão (ex: API morta no PM2, Connection Refused)
        logger.error(f"Erro de Conexão ao VERIFICAR número na porta {port}: {e}")
        raise WhatsAppInstanceDisconnectedException(f"API na porta {port} está offline (Connection Refused).")